#!/usr/bin/env python
'''
Author: Chris Duffy
ate: May 14, 2014
Name: docGenerator.py
Purpose: An script that can process dictionaries and lists and convert them into xlsx and docx files.
Input Data: "xml" XML file name that was passed by the parsing engine, used as a default name if no filename is passed
Input Data: "filename" filename that will be used to write other files
Input Data: "vulnerabilities" A dictionary of vulnerabilities{vulnerability id : [vulnerability title, vulnerability severity,vulnerability pciseverity, Vulnerability CVSS Score,Vulnerability CVSS Vector, Vulnerability Published Date, Vulnerability Modified Date, Vulnerability Updated Date, Vulnerability Description, References{Ref Type:Reference}, Solutions, Solution link, Vulnerability Status based on how it was identified]}
Input Data: "vuln_hosts" A dictionary of Vulnerabilities mapped to hosts vuln_hosts{Vulnerability IDs = [IPs, [hostname1, hostname2]]}
Input Data: "host_vulns" A dictionary of hosts mapped to details host_vulns{IPs=[MAC Addresses, [Hostnames], [Vulnerability IDs]]}
Input Data: "hosts" A dictionary of hosts{iterated number = [[Hostnames], IP address, protocol, port, service name, MAC Address]}
Input Data: "host_details" A dictionary of hosts mapped to details host_details{IPs=[MAC Addresses, [Hostnames], [Ports], [Services], [Port:Protocol], [Port:Protocol:Service], [Vulnerability IDs], Operating System Vendor, Operating System Product, Operating System Product]}
Input Data: "service_dict" A dictionary of services mapped to hosts service_dict{service=[IPs, [hostnames], ip:(hostname1, hostname2)]}
Input Data: "vuln_dict" A dictionary of vulnerability IDs mapped to vulnerability statuses vuln_dict{Vulnerability ID=Vulnerability Status} 
Input Data: "docVar" The type of office product that should be generated.

Copyright (c) 2015, Christopher Duffy All rights reserved.

Redistribution and use in source and binary forms, with or without modification, 
are permitted provided that the following conditions are met: * Redistributions 
of source code must retain the above copyright notice, this list of conditions and 
the following disclaimer. * Redistributions in binary form must reproduce the above 
copyright notice, this list of conditions and the following disclaimer in the 
documentation and/or other materials provided with the distribution. * Neither the 
name of the nor the names of its contributors may be used to endorse or promote 
products derived from this software without specific prior written permission.

THIS SOFTWARE IS PROVIDED BY THE COPYRIGHT HOLDERS AND CONTRIBUTORS 
"AS IS" AND ANY EXPRESS OR IMPLIED WARRANTIES, INCLUDING, BUT NOT LIMITED TO, 
THE IMPLIED WARRANTIES OF MERCHANTABILITY AND FITNESS FOR A PARTICULAR PURPOSE ARE 
DISCLAIMED. IN NO EVENT SHALL CHRISTOPHER DUFFY BE LIABLE FOR ANY DIRECT, INDIRECT, 
INCIDENTAL, SPECIAL, EXEMPLARY, OR CONSEQUENTIAL DAMAGES (INCLUDING, BUT NOT LIMITED TO, 
PROCUREMENT OF SUBSTITUTE GOODS OR SERVICES; LOSS OF USE, DATA, OR PROFITS; OR BUSINESS 
INTERRUPTION) HOWEVER CAUSED AND ON ANY THEORY OF LIABILITY, WHETHER IN CONTRACT, STRICT 
LIABILITY, OR TORT (INCLUDING NEGLIGENCE OR OTHERWISE) ARISING IN ANY WAY OUT OF THE USE 
OF THIS SOFTWARE, EVEN IF ADVISED OF THE POSSIBILITY OF SUCH DAMAGE.
'''
try:
    import docx
    from docx.shared import Inches
except:
    sys.exit("[!] Install the docx writer library as root or through sudo: pip install python-docx")
try:
    import xlsxwriter
except:
    sys.exit("[!] Install the xlsx writer library as root or through sudo: pip install xlsxwriter")
try:
    import pycurl
except:
    sys.exit("[!] Install the pycurl library as root or through sudo: pip install pycurl")

class docGenerator():
    def __init__(self, verbose, xml, filename, vulnerabilities, vuln_hosts, host_vulns, hosts, host_details, service_dict, vuln_dict, docVar):
        self.xml = xml
        self.filename = filename
        self.vulnerabilities = vulnerabilities
        self.vuln_hosts = vuln_hosts
        self.host_vulns = host_vulns
        self.hosts = hosts
        self.host_details = host_details
        self.service_dict = service_dict
        self.vuln_dict = vuln_dict
        self.docVar = docVar
        self.verbose = verbose
        try:
            self.run(verbose, xml, filename, vulnerabilities, vuln_hosts, host_vulns, hosts, host_details, service_dict, vuln_dict, docVar)
        except Exception as e:
            print(e)
    
    def run(self, verbose, xml, filename, vulnerabilities, vuln_hosts, host_vulns, hosts, host_details, service_dict, vuln_dict, docVar):
        print ("") #DEBUG
        # Run the appropriate module
        if docVar == "xlsx":
            if verbose > 0:
                print ("[*] Building %s.%s") % (filename, docVar)
            self.generateXSLX(verbose, xml, filename, vulnerabilities, vuln_hosts, host_vulns, hosts, host_details, service_dict, vuln_dict)
        elif docVar == "docx":
            self.generateDOCX(verbose, xml, filename, vulnerabilities, vuln_hosts, host_vulns, hosts, host_details, service_dict, vuln_dict)
        elif docVar == "all":
            self.generateDOCX(verbose, xml, filename, vulnerabilities, vuln_hosts, host_vulns, hosts, host_details, service_dict, vuln_dict)
            self.generateXSLX(verbose, xml, filename, vulnerabilities, vuln_hosts, host_vulns, hosts, host_details, service_dict, vuln_dict)
        else:
            sys.exit("[!] No document type requested")

    def curlModule(self, verbose, refDict):
        # Curl for Metasploit modules based on a reference passed
        # Input: Verbosity and reference list
        # Returned: exploit list
        exploitList=[]
        uniqRefsList=[]
        tempRefList=[]
        for k,v in refDict.items():
            if "CVE" in k:
                temp_ref = v
                tempRefList.append(temp_ref)
            elif "BID" in k:
                temp_ref="BID %s" % (v)
                tempRefList.append(temp_ref)
            elif "OSVDB" in k:
                temp_ref="OSVDB %s" % (v)
                tempRefList.append(temp_ref)
            else:
                pass
        if tempRefList is not None:
            uniqRefsList=uniqList(verbose, tempRefList)
            for ref in uniqRefsList:
                ref_encode=urllib.urlencode({'q':ref})
                query = "http://www.rapid7.com/db/search?utf8=%E2%9C%93&"+ref_encode+"&t=m"
                storage = StringIO()
                c = pycurl.Curl()
                c.setopt(c.URL, query)
                c.setopt(c.WRITEFUNCTION, storage.write)
                c.perform()
                c.close()
                content = storage.getvalue()
        else:
            exploitList.append("No Exploit Was Found")
        return (exploitList)

    def generateXSLX(self, verbose, xml, filename, vulnerabilities, vuln_hosts, host_vulns, hosts, host_details, service_dict, vuln_dict):
        references=[]
        exploit_temp=['No Exploits Found']
        if not filename:
            filename = "%s.xlsx" % (xml)
        else:
            filename = "%s.xlsx" % (filename)
        workbook = xlsxwriter.Workbook(filename)
        # Row one formatting
        format1 = workbook.add_format({'bold': True})
        format1.set_bg_color('#538DD5')
        # Even row formatting
        format2 = workbook.add_format({'text_wrap': True})
        format2.set_align('left')
        format2.set_align('top')
        format2.set_border(1)
        # Odd row formatting
        format3 = workbook.add_format({'text_wrap': True})
        format3.set_align('left')
        format3.set_align('top')
        format3.set_bg_color('#C5D9F1')
        format3.set_border(1)
        if verbose > 0:
            print ("[*] Creating Workbook: %s") % (filename)

        # Generate Worksheet 1
        worksheet = workbook.add_worksheet("Vulnerabilities")
        # Column width for worksheet 1
        worksheet.set_column(0, 0, 20)
        worksheet.set_column(1, 3, 12)
        worksheet.set_column(4, 5, 30)
        worksheet.set_column(6, 8, 19)
        worksheet.set_column(9, 9, 30)
        worksheet.set_column(10, 10, 13)
        worksheet.set_column(11, 12, 24)
        # Define starting location for Worksheet one
        row = 1
        col = 0
        # Generate Row 1 for worksheet one
        worksheet.write('A1', "Vulnerability Title", format1)
        worksheet.write('B1', "Severity", format1)
        worksheet.write('C1', "PCI Severity", format1)
        worksheet.write('D1', "CVSS Score", format1)
        worksheet.write('E1', "CVSS Vector", format1)
        worksheet.write('F1', "Affected Hosts", format1)
        worksheet.write('G1', "Published Date", format1)
        worksheet.write('H1', "Added Date", format1)
        worksheet.write('I1', "Modified Date", format1)
        worksheet.write('J1', "Description", format1)
        worksheet.write('K1', "References", format1)
        worksheet.write('L1', "Solutions", format1)
        worksheet.write('M1', "Solution Link", format1)
        worksheet.autofilter('A1:M1')
        # Populate Worksheet 1
        for key, value in vuln_hosts.items():        
            temp = str(vuln_dict.get(key)).strip('[]')
            if "not-vulnerable" not in temp:
                try:
                    temp=vulnerabilities[key]
                    vuln_title=temp[0]
                    vuln_severity=temp[1]
                    vuln_pciseverity=temp[2]
                    vuln_cvssscore=temp[3]
                    vuln_cvssvector=temp[4]
                    vuln_published=temp[5]
                    vuln_added=temp[6]
                    vuln_modified=temp[7]
                    vuln_description=temp[8]
                    ref_dict_temp_ws1=temp[9]
                    solutions=temp[10]
                    solution_link=temp[11]
                except:
                    if verbose > 3:
                        print "[!] An error occurred parsing vulnerbility ID: %s for Worksheet 1" %(key)
                hosts_temp = ",".join(vuln_hosts.get(key))
                hosts_temp = hosts_temp.split(':')
                hostnames_temp = str(hosts_temp[1]).strip('[]')
                hostnames_temp = "%s:(%s)"% (hosts_temp[0],hostnames_temp)
                for k, v in ref_dict_temp_ws1.items():
                    temps="%s:%s" % (k,v)
                    references.append(temps)
                ref = ", ".join(references)
                solutions = "".join(solutions)
                try:
                    if row % 2 != 0:
                        temp_format = format2
                    else:
                        temp_format = format3
                    if ref is None or ref == "":
                        ref="No References Supplied"
                    worksheet.write(row, col,     vuln_title, temp_format)
                    worksheet.write(row, col + 1, int(vuln_severity), temp_format)
                    worksheet.write(row, col + 2, int(vuln_pciseverity), temp_format)
                    worksheet.write(row, col + 3, float(vuln_cvssscore), temp_format)
                    worksheet.write(row, col + 4, vuln_cvssvector, temp_format)
                    worksheet.write(row, col + 5, hostnames_temp, temp_format)
                    worksheet.write(row, col + 6, vuln_published, temp_format)
                    worksheet.write(row, col + 7, vuln_added, temp_format)
                    worksheet.write(row, col + 8, vuln_modified, temp_format)
                    worksheet.write(row, col + 9, vuln_description, temp_format)
                    worksheet.write(row, col + 10, ref, temp_format)
                    worksheet.write(row, col + 11, solutions, temp_format)
                    worksheet.write(row, col + 12, solution_link, temp_format)
                    row += 1
                    references=[]
                except:
                    if verbose > 3:
                        print "[!] An error occurred writing data for vulnerability title: %s for Worksheet 1" % (vuln_title)

        # Generate Worksheet 2
        worksheet2 = workbook.add_worksheet("Hosts")
        # Column width for worksheet 2
        worksheet2.set_column(0, 0, 15)
        worksheet2.set_column(1, 1, 18)
        worksheet2.set_column(2, 2, 30)
        worksheet2.set_column(3, 4, 13)
        worksheet2.set_column(5, 5, 22)
        worksheet2.set_column(6, 9, 16)
        # Define starting location for Worksheet two
        row2 = 1
        col2 = 0 
        # Generate Row 1 for Worksheet two
        worksheet2.write('A1', "IP", format1)
        worksheet2.write('B1', "MAC Address", format1)
        worksheet2.write('C1', "Hostnames", format1)
        worksheet2.write('D1', "Ports", format1)
        worksheet2.write('E1', "Services", format1)
        worksheet2.write('F1', "Port:Protocol:Service", format1)
        worksheet2.write('G1', "Vulnerabilities", format1)
        worksheet2.write('H1', "Vendor", format1)
        worksheet2.write('I1', "Product", format1)
        worksheet2.write('J1', "Version", format1)
        worksheet2.autofilter('A1:J1')
        # Populate Worksheet 2
        for key, value in host_details.items():
            ip=key
            hwaddress=value[0]
            hostnames=", ".join(value[1])
            port_list=", ".join(value[2])
            service_list=", ".join(value[3])
            port_protocol_service_list=", ".join(value[5])
            num_vulns=value[6]
            os_vendor=value[7]
            os_product=value[8]
            os_version=value[9]
            try:
                if row2 % 2 != 0:
                    temp_format = format2
                else:
                    temp_format = format3
                worksheet2.write(row2, col2    , ip, temp_format)
                worksheet2.write(row2, col2 + 1, hwaddress, temp_format)
                worksheet2.write(row2, col2 + 2, hostnames, temp_format)
                worksheet2.write(row2, col2 + 3, port_list, temp_format)
                worksheet2.write(row2, col2 + 4, service_list, temp_format)
                worksheet2.write(row2, col2 + 5, port_protocol_service_list, temp_format)
                worksheet2.write(row2, col2 + 6, int(num_vulns), temp_format)
                worksheet2.write(row2, col2 + 7, os_vendor, temp_format)
                worksheet2.write(row2, col2 + 8, os_product, temp_format)
                worksheet2.write(row2, col2 + 9, os_version, temp_format)
            except:
                if verbose > 3:
                    print "[!] An error occurred writing data for IP: %s in Worksheet 2" % (ip)
            row2 += 1

        # Write Worksheet 3
        worksheet3 = workbook.add_worksheet("Services")
        # Column width for worksheet 3
        worksheet3.set_column(0, 1, 30)
        # Define starting location for Worksheet three
        row3 = 1
        col3 = 0
        # Generate Row 1 for Worksheet three
        worksheet3.write('A1', "Service", format1)
        worksheet3.write('B1', "Hosts", format1)
        worksheet3.autofilter('A1:B1')
        # Populate Worksheet 3
        for key, value in service_dict.items():
            service=key   
            host = ", ".join(value[0])
            try:
                if row3 % 2 != 0:
                    temp_format = format2
                else:
                    temp_format = format3
                worksheet3.write(row3, col3,     service, temp_format)
                worksheet3.write(row3, col3 + 1, host, temp_format)
            except:
                if verbose > 3:
                    print "[!] An error occurred writing data for service: %s in Worksheet 3" % (service)
            row3 += 1

        # Generate Worksheet 4
        worksheet4 = workbook.add_worksheet("MiTM Targets")
        # Column width for worksheet 4
        worksheet4.set_column(0, 0, 15)
        worksheet4.set_column(1, 1, 18)
        worksheet4.set_column(2, 2, 30)
        worksheet4.set_column(3, 3, 22)
        # Define starting location for Worksheet four
        row4 = 1
        col4 = 0
        # Generate Row 1 for worksheet four
        worksheet4.write('A1', "IP", format1)
        worksheet4.write('B1', "MAC Address", format1)
        worksheet4.write('C1', "Hostnames", format1)
        worksheet4.write('D1', "Port:Protocol:Service", format1)
        worksheet4.autofilter('A1:D1')
        # Populate Worksheet 4
        for key, value in host_details.items():
            ip=key
            hwaddress=value[0]
            hostnames=", ".join(value[1])
            port_protocol_service_list=", ".join(value[5])
            if "Undiscovered" not in hwaddress:
                try:
                    if row4 % 2 != 0:
                        temp_format = format2
                    else:
                        temp_format = format3
                    worksheet4.write(row4, col4    , ip, temp_format)
                    worksheet4.write(row4, col4 + 1, hwaddress, temp_format)
                    worksheet4.write(row4, col4 + 2, hostnames, temp_format)
                    worksheet4.write(row4, col4 + 3, port_protocol_service_list, temp_format)
                except:
                    if verbose > 3:
                        print "[!] An error occurred writing data for IP: %s in Worksheet 4" % (ip)
                row4 += 1

        # Generate Worksheet 5
        worksheet5 = workbook.add_worksheet("Exploitable Targets")
        # Column width for worksheet 5
        worksheet5.set_column(0, 0, 30)
        worksheet5.set_column(1, 1, 25)
        worksheet5.set_column(2, 2, 30)
        worksheet5.set_column(3, 3, 19)
        worksheet5.set_column(4, 4, 20)
        # Define starting location for Worksheet five
        row5 = 1
        col5 = 0
        # Generate Row 1 for worksheet five
        worksheet5.write('A1', "Vulnerability Title", format1)
        worksheet5.write('B1', "Affected Hosts", format1)
        worksheet5.write('C1', "Description", format1)
        worksheet5.write('D1', "Exploit", format1)
        worksheet5.write('E1', "References", format1)
        worksheet5.autofilter('A1:E1')
        # Populate Worksheet 5
        for key, value in vuln_hosts.items():
            temp = str(vuln_dict.get(key)).strip('[]')
            if "exploit" in temp:
                try:
                    temp=vulnerabilities[key]
                    vuln_title=temp[0]
                    vuln_description=temp[8]
                    ref_dict_temp=temp[9]
                except:
                    if verbose > 3:
                        print "[!] An error occurred parsing vulnerbility ID: %s  for Worksheet 5" %(key)
                hosts_temp = ",".join(vuln_hosts.get(key))
                hosts_temp = hosts_temp.split(':')
                hostnames_temp = str(hosts_temp[1]).strip('[]')
                hostnames_temp = "%s:(%s)"% (hosts_temp[0],hostnames_temp)
                #exploit_temp = curlModule(verbose, ref_dict_temp)
                exploits = ", ".join(exploit_temp)
                for k, v in ref_dict_temp.items():
                    temps="%s:%s" % (k,v)
                    references.append(temps)
                ref = ", ".join(references)
                if ref is None or ref == "":
                    ref="No References Supplied"
                try:
                    if row5 % 2 != 0:
                        temp_format = format2
                    else:
                        temp_format = format3
                    worksheet5.write(row5, col5,     vuln_title, temp_format)
                    worksheet5.write(row5, col5 + 1, hostnames_temp, temp_format)
                    worksheet5.write(row5, col5 + 2, vuln_description, temp_format)
                    worksheet5.write(row5, col5 + 3, exploits, temp_format)
                    worksheet5.write(row5, col5 + 4, ref, temp_format)
                    row5 += 1
                    references=[]
                except:
                    if verbose > 3:
                        print "[!] An error occurred writing data for vulnerability title: %s in Worksheet 5" % (vuln_title)

        # Generate Worksheet 6
        worksheet6 = workbook.add_worksheet("Vulnerable Software Versions")
        # Column width for worksheet 6
        worksheet6.set_column(0, 0, 30)
        worksheet6.set_column(1, 1, 25)
        worksheet6.set_column(2, 2, 30)
        worksheet6.set_column(3, 3, 20)
        # Define starting location for Worksheet six
        row6 = 1
        col6 = 0
        # Generate Row 1 for worksheet six
        worksheet6.write('A1', "Vulnerability Title", format1)
        worksheet6.write('B1', "Affected Hosts", format1)
        worksheet6.write('C1', "Description", format1)
        worksheet6.write('D1', "References", format1)
        worksheet6.autofilter('A1:D1')
        # Populate Worksheet 6
        for key, value in vuln_hosts.items():
            temp = str(vuln_dict.get(key)).strip('[]')
            if "vulnerable-version" in temp:
                try:
                    temp=vulnerabilities[key]
                    vuln_title=temp[0]
                    vuln_description=temp[8]
                    ref_dict_temp=temp[9]
                except:
                    if verbose > 3:
                        print "[!] An error occurred parsing vulnerbility ID: %s in Workseet 6" %(key)
                hosts_temp = ",".join(vuln_hosts.get(key))
                hosts_temp = hosts_temp.split(':')
                hostnames_temp = str(hosts_temp[1]).strip('[]')
                hostnames_temp = "%s:(%s)"% (hosts_temp[0],hostnames_temp)
                for k, v in ref_dict_temp.items():
                    temps="%s:%s" % (k,v)
                    references.append(temps)
                ref = ", ".join(references)
                if ref is None:
                    ref="No References Supplied"
                try:
                    if row6 % 2 != 0:
                        temp_format = format2
                    else:
                        temp_format = format3
                    worksheet6.write(row6, col6,     vuln_title, temp_format)
                    worksheet6.write(row6, col6 + 1, hostnames_temp, temp_format)
                    worksheet6.write(row6, col6 + 2, vuln_description, temp_format)
                    worksheet6.write(row6, col6 + 3, ref, temp_format)
                    row6 += 1
                    references=[]
                except:
                    if verbose > 3:
                        print "[!] An error occurred writing data for vulnerability title: %s in Worksheet 6" % (vuln_title)

        # Generate Worksheet 7
        worksheet7 = workbook.add_worksheet("IP to Vulnerability Matrix")
        # Column width for worksheet 7
        worksheet7.set_column(0, 0, 19)
        worksheet7.set_column(1, 1, 20)
        worksheet7.set_column(2, 3, 12)
        worksheet7.set_column(4, 5, 30)
        worksheet7.set_column(7, 8, 19)
        worksheet7.set_column(9, 9, 30)
        worksheet7.set_column(10, 10, 13)
        worksheet7.set_column(11, 12, 24)
        # Define starting location for Worksheet seven
        row7 = 1
        col7 = 0
        # Generate Row 1 for worksheet 7
        worksheet7.write('A1', "Affected Host", format1)
        worksheet7.write('B1', "Vulnerability Title", format1)
        worksheet7.write('C1', "Severity", format1)
        worksheet7.write('D1', "PCI Severity", format1)
        worksheet7.write('E1', "CVSS Score", format1)
        worksheet7.write('F1', "CVSS Vector", format1)
        worksheet7.write('G1', "Published Date", format1)
        worksheet7.write('H1', "Added Date", format1)
        worksheet7.write('I1', "Modified Date", format1)
        worksheet7.write('J1', "Description", format1)
        worksheet7.write('K1', "References", format1)
        worksheet7.write('L1', "Solutions", format1)
        worksheet7.write('M1', "Solution Link", format1)
        worksheet7.autofilter('A1:M1')
        # Populate Worksheet 7
        for key, value in vuln_hosts.items():        
            try:
                temp=vulnerabilities[key]
                vuln_title=temp[0]
                vuln_severity=temp[1]
                vuln_pciseverity=temp[2]
                vuln_cvssscore=temp[3]
                vuln_cvssvector=temp[4]
                vuln_published=temp[5]
                vuln_added=temp[6]
                vuln_modified=temp[7]
                vuln_description=temp[8]
                ref_dict_temp_ws1=temp[9]
                solutions=temp[10]
                solution_link=temp[11]
            except:
                if verbose > 3:
                    print "[!] An error occurred parsing vulnerbility ID: %s for Worksheet 7" %(key)
            for host in vuln_hosts.get(key):
                host = host.split(':')
                hostnames_temp = str(host[1]).strip('[]')
                hostnames_temp = "%s:(%s)"% (host[0],hostnames_temp)
            for k, v in ref_dict_temp_ws1.items():
                temps="%s:%s" % (k,v)
                references.append(temps)
            ref = ", ".join(references)
            solutions = "".join(solutions)
            try:
                if row7 % 2 != 0:
                    temp_format = format2
                else:
                    temp_format = format3
                if ref is None or ref == "":
                    ref="No References Supplied"
                worksheet7.write(row7, col7,     hostnames_temp, temp_format)
                worksheet7.write(row7, col7 + 1, vuln_title, temp_format)
                worksheet7.write(row7, col7 + 2, int(vuln_severity), temp_format)
                worksheet7.write(row7, col7 + 3, int(vuln_pciseverity), temp_format)
                worksheet7.write(row7, col7 + 4, float(vuln_cvssscore), temp_format)
                worksheet7.write(row7, col7 + 5, vuln_cvssvector, temp_format)
                worksheet7.write(row7, col7 + 6, vuln_published, temp_format)
                worksheet7.write(row7, col7 + 7, vuln_added, temp_format)
                worksheet7.write(row7, col7 + 8, vuln_modified, temp_format)
                worksheet7.write(row7, col7 + 9, vuln_description, temp_format)
                worksheet7.write(row7, col7 + 10, ref, temp_format)
                worksheet7.write(row7, col7 + 11, solutions, temp_format)
                worksheet7.write(row7, col7 + 12, solution_link, temp_format)
                row7 += 1
                references=[]
            except:
                if verbose > 3:
                    print "[!] An error occurred writing data for host: %s in Worksheet 7" % (host)

        # Generate Worksheet 8
        worksheet8 = workbook.add_worksheet("Not Vulnerable or FP")
        # Column width for worksheet 8
        worksheet8.set_column(0, 0, 30)
        worksheet8.set_column(1, 1, 25)
        worksheet8.set_column(2, 2, 30)
        worksheet8.set_column(3, 3, 20)
        worksheet8.set_column(4, 4, 30)
        # Define starting location for Worksheet eight
        row8 = 1
        col8 = 0
        # Generate Row 1 for worksheet eight
        worksheet8.write('A1', "Vulnerability Title", format1)
        worksheet8.write('B1', "Affected Hosts", format1)
        worksheet8.write('C1', "Description", format1)
        worksheet8.write('D1', "References", format1)
        worksheet8.write('E1', "Results", format1)
        worksheet8.autofilter('A1:E1')
        # Populate Worksheet 8
        for key, value in vuln_hosts.items():
            temp = str(vuln_dict.get(key)).strip('[]')
            if "not-vulnerable" in temp:
                try:
                    temp=vulnerabilities[key]
                    vuln_title=temp[0]
                    vuln_description=temp[8]
                    ref_dict_temp=temp[9]
                except:
                    if verbose > 3:
                        print "[!] An error occurred parsing vulnerbility ID: %s in Worksheet 8" %(key)
                hosts_temp = ",".join(vuln_hosts.get(key))
                hosts_temp = hosts_temp.split(':')
                hostnames_temp = str(hosts_temp[1]).strip('[]')
                hostnames_temp = "%s:(%s)"% (hosts_temp[0],hostnames_temp)
                #exploit_temp = self.curlModule(verbose, ref_dict_temp)
                exploits = ", ".join(exploit_temp)
                for k, v in ref_dict_temp.items():
                    temps="%s:%s" % (k,v)
                    references.append(temps)
                ref = ", ".join(references)
                if ref is None or ref == "":
                    ref="No References Supplied"
                try:
                    if row8 % 2 != 0:
                        temp_format = format2
                    else:
                        temp_format = format3
                    worksheet8.write(row8, col8,     vuln_title, temp_format)
                    worksheet8.write(row8, col8 + 1, hostnames_temp, temp_format)
                    worksheet8.write(row8, col8 + 2, vuln_description, temp_format)
                    worksheet8.write(row8, col8 + 3, ref, temp_format)
                    worksheet8.write(row8, col8 + 4, "Not Vulnerable or False Positive", temp_format)
                    row8 += 1
                    references=[]
                except:
                    if verbose > 3:
                        print "[!] An error occurred writing data for vulnerability title: %s in Worksheet 8" % (vuln_title)

        # Generate worksheet 9
        worksheet9 = workbook.add_worksheet("IP to Service Itemization")
        # Column width for worksheet 9
        worksheet9.set_column(0, 0, 15)
        worksheet9.set_column(1, 1, 18)
        worksheet9.set_column(2, 2, 30)
        worksheet9.set_column(3, 4, 13)
        worksheet9.set_column(5, 5, 13)
        worksheet9.set_column(6, 6, 22)
        # Define starting location for Worksheet nine
        row9 = 1
        col9 = 0
        # Generate Row 1 for worksheet nine
        worksheet9.write('A1', "IP", format1)
        worksheet9.write('B1', "MAC Address", format1)
        worksheet9.write('C1', "Hostnames", format1)
        worksheet9.write('D1', "Ports", format1)
        worksheet9.write('E1', "Protocol", format1)
        worksheet9.write('F1', "Services", format1)
        worksheet9.autofilter('A1:F1')
        # Populate worksheet 9
        for key, value in hosts.items():
            hostnames=value[0]
            ip=value[1]
            proto=value[2]
            port=value[3]
            service=value[4]
            hwaddress=value[5]
            hostnames = ",".join(hostnames)
            try:
                if row9 % 2 != 0:
                    temp_format = format2
                else:
                    temp_format = format3
                worksheet9.write(row9, col9    , ip, temp_format)
                worksheet9.write(row9, col9 + 1, hwaddress, temp_format)
                worksheet9.write(row9, col9 + 2, hostnames, temp_format)
                worksheet9.write(row9, col9 + 3, int(port), temp_format)
                worksheet9.write(row9, col9 + 4, proto, temp_format)
                worksheet9.write(row9, col9 + 5, service, temp_format)
            except:
                if verbose > 3:
                    print "[!] An error occurred writing data for IP: %s in Worksheet 9" % (ip)
            row9 += 1

        try:
            workbook.close()
        except:
            sys.exit("[!] Permission to write to the file or location provided was denied")

    def generateDOCX(self, verbose, xml, filename, vulnerabilities, vuln_hosts, host_vulns, hosts, host_details, service_dict, vuln_dict):
        references=[]
        temp_list=[]
        high_to_low=[]
        host_list=[]
        document = docx.Document()
        crit_count=1
        high_count=1
        med_count=1
        low_count=1
        info_count=1
        if not filename:
            filename = "%s.docx" % (xml)
        else:
            filename = "%s.docx" % (filename)
        for key, value in vuln_hosts.items():        
            temp = str(vuln_dict.get(key)).strip('[]')
            if "not-vulnerable" not in temp:
                try:
                    temp=vulnerabilities[key]
                    vuln_title=temp[0]
                    vuln_severity=temp[1]
                    vuln_pciseverity=temp[2]
                    vuln_cvssscore=temp[3]
                    vuln_cvssvector=temp[4]
                    vuln_published=temp[5]
                    vuln_added=temp[6]
                    vuln_modified=temp[7]
                    vuln_description=temp[8]
                    ref_dict_temp_ws1=temp[9]
                    solutions=temp[10]
                    solution_link=temp[11]
                except:
                    if verbose > 3:
                        print "[!] An error occurred parsing vulnerbility ID: %s for Document: %s" %(key, filename)
                hosts_temp = ",".join(vuln_hosts.get(key))
                hosts_temp = hosts_temp.split(':')
                hostnames_temp = str(hosts_temp[1]).strip('[]')
                hostnames_temp = ("%s:(%s)"% (hosts_temp[0],hostnames_temp))
                host_list.append(hostnames_temp)
                for k, v in ref_dict_temp_ws1.items():
                    temps="%s:%s" % (k,v)
                    references.append(temps)
                ref = ", ".join(references)
                solutions = "".join(solutions)
                temp_tup = (vuln_title, vuln_severity, vuln_pciseverity, vuln_cvssscore, vuln_cvssvector, vuln_published, vuln_added, vuln_modified, vuln_description, ref, solutions, solution_link, host_list)
                temp_list.append(temp_tup)
                references=[]
                host_list=[]
        high_to_low=sorted(temp_list, key=lambda x: float(x[3]), reverse=True)
        for item in high_to_low:
            vuln_title=item[0]
            vuln_severity=item[1]
            vuln_pciseverity=item[2]
            vuln_cvssscore=item[3]
            vuln_cvssvector=item[4]
            vuln_published=item[5]
            vuln_added=item[6]
            vuln_modified=item[7]
            vuln_description=item[8]
            ref=item[9]
            solutions=item[10]
            solution_link=item[11]
            hosts=item[12]
            if float(vuln_cvssscore) >= 7.5:
                if crit_count ==1:
                    document.add_heading('Critical Findings', level=2)
                finding_name=("6.1.%s %s") % (crit_count, vuln_title)
                crit_count +=1
            elif float(vuln_cvssscore) > 5 and float(vuln_cvssscore) < 7.5:
                if high_count ==1:
                    document.add_heading('High Findings', level=2)
                finding_name=("6.2.%s %s") % (high_count, vuln_title)
                high_count +=1
            elif float(vuln_cvssscore) > 2.5 and float(vuln_cvssscore) < 5.1 :
                if med_count ==1:
                    document.add_heading('Medium Findings', level=2)
                finding_name=("6.3.%s %s") % (med_count, vuln_title)
                med_count +=1
            elif float(vuln_cvssscore) <= 2.5:
                if low_count ==1:
                    document.add_heading('Low Findings', level=2)
                finding_name=("6.4.%s %s") % (low_count, vuln_title)
                low_count +=1
            else:
                if info_count==1:
                    document.add_heading('Informational Findings', level=2)
                finding_name=("6.5.%s %s") % (info_count, vuln_title)
                info_count +=1
            document.add_heading(finding_name, level=3)
            table = document.add_table(rows=1,cols=7)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text='CVSS Base Score'
            hdr_cells[1].text='Impact Sub-score'
            hdr_cells[2].text='Exploitability Sub-score'
            hdr_cells[3].text='CVSS Temporal Score'
            hdr_cells[4].text='CVSS Enviromental Score'
            hdr_cells[5].text='Modified Impact Sub-score'
            hdr_cells[6].text='Overall CVSS Score'
            row_cells = table.add_row().cells
            row_cells[0].text = str(vuln_cvssscore)
            row_cells[1].text = str("N/A")
            row_cells[2].text = str("N/A")
            row_cells[3].text = str("N/A")
            row_cells[4].text = str("N/A")
            row_cells[5].text = str("N/A")
            row_cells[6].text = str("N/A")
            sub_head1 = document.add_paragraph()
            sub_head1.add_run('Ease of Exploitation:').bold=True
            sub_head2 = document.add_paragraph()
            sub_head2.add_run("Mitigation Level: ").bold=True
            sub_head3 = document.add_paragraph()
            sub_head3.add_run("Summary: ").bold=True
            finding_description = document.add_paragraph(vuln_description)
            sub_head4 = document.add_paragraph()
            sub_head4.add_run("Affected Host/Locations: ").bold=True
            for host in hosts:
                document.add_paragraph(host, style='ListBullet')
            sub_head5 = document.add_paragraph()
            sub_head5.add_run("Proof of Concept: ").bold=True
            sub_head6 = document.add_paragraph()
            sub_head6.add_run("Recommendation: ").bold=True
            sub_head8 = document.add_paragraph("Tactical Remediation: ")
            solutions_para = document.add_paragraph()
            solutions_para.add_run(solutions)
            solutions_link_para =document.add_paragraph()
            solutions_link_para.add_run(solution_link).italic=True
            sub_head9 = document.add_paragraph("Strategic Remediation: ")
            sub_head7 = document.add_paragraph()
            sub_head7.add_run("References: ").bold=True
            if ref is None or ref is "":
                ref="No References Supplied"
            finding_references = document.add_paragraph(ref)
            document.add_page_break()
        if verbose > 0:
            print ("[*] Creating Document: %s") % (filename)
        try:
            document.save(filename)
        except:
            sys.exit("[!] Permission to write to the file or location provided was denied")
